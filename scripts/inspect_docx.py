import docx

doc = docx.Document('Advanced_Detection_of_Pre-Ictal_Stages_in_Epilepsy(8.7.2026).docx')
print('=== Document Sections & Margins ===')
for i, sec in enumerate(doc.sections):
    print(f'Section {i}: top={sec.top_margin.pt}pt, bottom={sec.bottom_margin.pt}pt, left={sec.left_margin.pt}pt, right={sec.right_margin.pt}pt, page_width={sec.page_width.pt}pt, page_height={sec.page_height.pt}pt')

print('\n=== Paragraph Styles & Fonts Sample (First 35 paragraphs) ===')
for i, p in enumerate(doc.paragraphs[:35]):
    if p.text.strip():
        fonts = set(r.font.name for r in p.runs if r.font.name)
        sizes = set(r.font.size.pt for r in p.runs if r.font.size)
        colors = set(str(r.font.color.rgb) for r in p.runs if r.font.color and r.font.color.rgb)
        bolds = set(r.bold for r in p.runs)
        print(f'P{i:02d} [{p.style.name}] (align={p.alignment}): \"{p.text[:60]}...\" | fonts={fonts}, sizes={sizes}, colors={colors}, bold={bolds}')

print('\n=== Table Styles Sample ===')
for i, t in enumerate(doc.tables[:3]):
    print(f'Table {i}: style={t.style.name}, rows={len(t.rows)}, cols={len(t.columns)}')
    for r in t.rows[:3]:
        row_txt = [c.text.strip().replace('\n', ' ') for c in r.cells]
        print('   ', row_txt)
