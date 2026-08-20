import docx

doc = docx.Document('Advanced_Detection_of_Pre-Ictal_Stages_in_Epilepsy(8.7.2026).docx')
print('=== Style Font Details ===')
for s in doc.styles:
    if hasattr(s, 'font') and s.font and s.font.name:
        print(f'Style: {s.name}, Font: {s.font.name}, Size: {s.font.size.pt if s.font.size else None}')

print('\n=== Elements & Heading Inspection ===')
for p in doc.paragraphs:
    if 'Heading' in p.style.name or p.style.name in ['Title', 'Subtitle'] or (p.runs and p.runs[0].bold and len(p.text) < 80):
        print(f'Heading/Bold [{p.style.name}]: \"{p.text}\"')

print('\n=== References Section Inspection ===')
ref_started = False
for p in doc.paragraphs:
    if 'Reference' in p.text or 'REFERENCES' in p.text:
        ref_started = True
    if ref_started:
        print(f'Ref [{p.style.name}]: {p.text[:90]}...')
